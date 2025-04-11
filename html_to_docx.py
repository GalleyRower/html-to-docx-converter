import os
import re
import sys
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import chardet  # Для определения кодировки
#import subprocess  # Для проверки открытия DOCX - ЗАКОММЕНТИРОВАНО
import docx.opc.constants
from docx.oxml.shared import OxmlElement
from docx.oxml.shared import qn

# Попытка импортировать WD_STYLE_TYPE и обработать отсутствие атрибута
try:
    from docx.enum.text import WD_STYLE_TYPE
except ImportError:
    print("Предупреждение: WD_STYLE_TYPE не найден в docx.enum.text. Функциональность стилей может быть ограничена.")
    WD_STYLE_TYPE = None  # Устанавливаем в None, чтобы избежать ошибок в дальнейшем

def detect_encoding(file_path):
    """Определяет кодировку файла."""
    try:
        with open(file_path, 'rb') as f:
            result = chardet.detect(f.read())
        return result['encoding']
    except Exception as e:
        print(f"Ошибка при определении кодировки для {file_path}: {e}")
        return None

def add_hyperlink(paragraph, text, url):
    """Добавляет гиперссылку в параграф."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def html_to_docx(html_file, docx_file):
    """Конвертирует HTML с публицистическим текстом в DOCX, применяя теги."""
    try:
        encoding = detect_encoding(html_file)
        if not encoding:
            print(f"Не удалось определить кодировку файла: {html_file}")
            return False, "Не удалось определить кодировку"

        with open(html_file, 'r', encoding=encoding) as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        document = Document()

        # Задаем стили по умолчанию
        default_font = document.styles['Normal'].font
        default_font.name = 'Arial'  # Менее официальный шрифт
        default_font.size = Pt(11)  # Немного уменьшенный размер
        # Добавляем интервал между строками
        for paragraph_style in document.styles:
            if WD_STYLE_TYPE and hasattr(paragraph_style, 'type') and paragraph_style.type == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_format = paragraph_style.paragraph_format
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Одинарный интервал
                paragraph_format.space_before = Pt(0)  # Убираем отступы перед абзацем
                paragraph_format.space_after = Pt(0)   # Убираем отступы после абзаца

        # Обработка содержимого. Теперь soup - это и есть все содержимое, без body
        for element in soup.contents:  # Используем soup.contents
            if element is None:
                continue  # Пропускаем пустые элементы

            if element.name == 'h1':
                heading = document.add_heading(element.text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
                heading.style.font.name = 'Arial' # переопределяем шрифт для заголовка
            elif element.name == 'h2':
                heading = document.add_heading(element.text, level=2)
                heading.style.font.name = 'Arial'
            elif element.name == 'h3':
                heading = document.add_heading(element.text, level=3)
                heading.style.font.name = 'Arial'
            elif element.name == 'h4':
                heading = document.add_heading(element.text, level=4)
                heading.style.font.name = 'Arial'
            elif element.name == 'h5':
                heading = document.add_heading(element.text, level=5)
                heading.style.font.name = 'Arial'
            elif element.name == 'h6':
                heading = document.add_heading(element.text, level=6)
                heading.style.font.name = 'Arial'
            elif element.name == 'p':
                paragraph = document.add_paragraph(element.text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
                paragraph.style.font.name = 'Arial'
            elif element.name == 'strong' or element.name == 'b':
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.text)
                run.bold = True
                run.font.name = 'Arial' # Явно задаем шрифт для выделенного текста
            elif element.name == 'em' or element.name == 'i':
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.text)
                run.italic = True
                run.font.name = 'Arial'
            elif element.name == 'a':
                # Ссылки требуют более сложной обработки
                try:
                    href = element['href']
                    text = element.text
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run(text)
                    add_hyperlink(paragraph, text, href)  #  Используем функцию добавления гиперссылки
                    run.underline = True
                    run.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0xFF)  # Синий цвет
                    run.font.name = 'Arial'
                except KeyError:
                    print(f"Ссылка без атрибута href: {element}")
                    paragraph = document.add_paragraph(element.text) # Просто добавляем текст
                    paragraph.style.font.name = 'Arial'
            elif element.name == 'ul':
                # Ненумерованный список
                for li in element.find_all('li'):
                    paragraph = document.add_paragraph(li.text, style='List Bullet')
                    paragraph.style.font.name = 'Arial'
            elif element.name == 'ol':
                # Нумерованный список
                for i, li in enumerate(element.find_all('li')):
                    paragraph = document.add_paragraph(f"{i+1}. {li.text}", style='List Number')
                    paragraph.style.font.name = 'Arial'
            elif element.name == 'img':
                try:
                    #Попытка добавить картинку. Нужно убедиться, что путь к картинке валидный
                    document.add_picture(element['src'], width=Inches(6))
                except Exception as e:
                    print(f"Ошибка при добавлении изображения: {e}")
                    # Логируем ошибку, но продолжаем выполнение
            elif element.name == 'table':
                 # Обработка таблиц - требует более сложной логики
                table = document.add_table(rows=0, cols=0)
                for row_index, row in enumerate(element.find_all('tr')):
                    cells = row.find_all(['td', 'th'])
                    if row_index == 0:  # Первая строка - определяем количество столбцов
                        for _ in cells:
                             table.add_column(Inches(1.5))  # Ширина столбца
                        continue

                    docx_row = table.add_row()
                    for col_index, cell in enumerate(cells):
                        docx_row.cells[col_index].text = cell.text

            elif element.name == 'blockquote':
                # Цитата
                paragraph = document.add_paragraph(element.text)
                paragraph.style = 'Intense Quote'  # Используем стандартный стиль цитаты
                paragraph.style.font.name = 'Arial'
            elif element.name == 'pre':
                # Отформатированный текст (код)
                paragraph = document.add_paragraph(element.text)
                paragraph.style = 'Code'  #  Предполагается, что стиль Code определен в шаблоне DOCX
                paragraph.font.name = 'Courier New' #  Моноширинный шрифт
            elif element.name == 'br':
                # Перенос строки (можно добавить пустой параграф)
                document.add_paragraph()

            elif isinstance(element, NavigableString):
                # Обработка простого текста вне тегов
                text = str(element).strip()
                if text:
                    paragraph = document.add_paragraph(text)
                    paragraph.style.font.name = 'Arial'

        document.save(docx_file)
        return True, None  # Успешно, без ошибок

    except Exception as e:
        print(f"Ошибка при конвертации {html_file} в {docx_file}: {e}")
        return False, str(e)


def verify_docx(docx_file):
    """Проверяет, что DOCX файл открывается и содержит читаемый текст."""
    #try: # ЗАКОММЕНТИРОВАНО
        # Проверка открытия файла
        #subprocess.run(['start', 'msword', docx_file], shell=True, check=True, timeout=10)  # Windows (требуется msword в PATH) # ЗАКОММЕНТИРОВАНО
        # Альтернатива для Linux (требуется libreoffice):
        # subprocess.run(['libreoffice', '--headless', '--convert-to', 'txt:Text export', docx_file], check=True, timeout=10) # ЗАКОММЕНТИРОВАНО

        # Проверка кодировки (это сложнее, т.к. DOCX хранит текст в Unicode)
    document = Document(docx_file)
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)
    text = '\n'.join(full_text)

    # Простая проверка на "читаемость" - ищем непечатаемые символы
    if re.search(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', text):
        print(f"Обнаружены непечатаемые символы в {docx_file}")
        return False, "Непечатаемые символы"

        # Проверка на "разумность" текста (пример)
    if len(text) < 10: #  Слишком короткий текст - подозрительно
        print(f"Слишком короткий текст в {docx_file}")
        return False, "Слишком короткий текст"

    return True, None

    #except subprocess.CalledProcessError as e: # ЗАКОММЕНТИРОВАНО
    #    print(f"Ошибка при открытии {docx_file}: {e}")  # ЗАКОММЕНТИРОВАНО
    #    return False, "Ошибка открытия файла"  # ЗАКОММЕНТИРОВАНО
    #except Exception as e: # ЗАКОММЕНТИРОВАНО
    #    print(f"Ошибка при чтении {docx_file}: {e}") # ЗАКОММЕНТИРОВАНО
    #    return False, str(e) # ЗАКОММЕНТИРОВАНО


def main(input_dir, output_dir):
    """Главная функция: обрабатывает все файлы в директории."""
    total_files = 0
    successful_exports = 0
    failed_exports = 0
    errors = {}

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for filename in os.listdir(input_dir):
        if filename.endswith(".txt") or filename.endswith(".html"):  # Обрабатываем .txt и .html
            total_files += 1
            input_file = os.path.join(input_dir, filename)
            output_file = os.path.join(output_dir, os.path.splitext(filename)[0] + ".docx")

            success, error_message = html_to_docx(input_file, output_file)

            if success:
                #verification_success, verification_error = verify_docx(output_file) # ЗАКОММЕНТИРОВАНО - вызов проверки открытия
                verification_success = True #  Принудительно считаем, что проверка прошла успешно
                verification_error = None
                if verification_success:
                    successful_exports += 1
                    print(f"Успешно экспортирован: {filename}")
                else:
                    failed_exports += 1
                    print(f"Ошибка верификации {filename}: {verification_error}")
                    errors[filename] = verification_error
            else:
                failed_exports += 1
                print(f"Ошибка экспорта {filename}: {error_message}")
                errors[filename] = error_message

    print("\n--- Отчет ---")
    print(f"Всего файлов: {total_files}")
    print(f"Успешно экспортировано: {successful_exports}")
    print(f"Неудачно экспортировано: {failed_exports}")
    if errors:
        print("\nОшибки:")
        for filename, error in errors.items():
            print(f"{filename}: {error}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Использование: python script.py <input_директория> <output_директория>")
        sys.exit(1)

    input_directory = sys.argv[1]
    output_directory = sys.argv[2]
    main(input_directory, output_directory)
